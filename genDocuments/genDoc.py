from docx import Document
from docx.shared import Pt
import json
import os

version = "a0.0.1"


def mkdir(path):

    folder = os.path.exists(path)

    if not folder:
        os.makedirs(path)
    else:
        os.path.exists(path)


if __name__ == "__main__":

    copyright_content = "Generated By HAM Examination Paper Generator Developed By Yurik"

    contest_id = input("contest_id: ")

    if os.path.exists(contest_id):

        mkdir("./" + contest_id + "/Documents/")

        #  Create Test Paper

        with open("./"+contest_id+"/Paper.json", "r", encoding="utf-8") as f:
            contest_que = json.load(f)

        document = Document()

        document.add_heading(
            '业余无线电模拟试题'+str(contest_que['contestId'])+"_"+str(contest_que['contestGroup']))
        document.add_paragraph(
            '一、单选题 (共'+str(len(contest_que['questions']))+'小题)')

        count = 1

        for i in contest_que['questions']:
            # print(count)
            document.add_paragraph(str(count)+". "+i['question'])
            count = count+1
            document.add_paragraph(
                "A."+i['A']+"\n"+"B."+i['B']+"\n"+"C."+i['C']+"\n"+"D."+i['D']+"\n")

        document.add_paragraph(
            copyright_content)

        document.save("./"+contest_id+"/Documents/Paper.docx")

        #  Create Answers Paper

        with open("./"+contest_id+"/Answer.json", "r", encoding="utf-8") as f:
            contest_que = json.load(f)

        document = Document()

        document.add_heading(
            '业余无线电模拟试题'+str(contest_que['contestId'])+"_"+str(contest_que['contestGroup'])+"答案")

        document.add_paragraph(
            '一、单选题 (共'+str(len(contest_que['questions']))+'小题)')

        count = 0

        quick_check = document.add_paragraph()

        for i in contest_que['questions']:
            if (count % 5 == 0):
                quick_check.add_run("   ")

            quick_check.add_run(i['option'])
            count = count+1

        document.add_paragraph()

        count = 1

        for i in contest_que['questions']:
            document.add_paragraph(
                str(count)+'. (LK'+str(i['uid']).zfill(4)+') '+i['question'])
            count = count+1
            document.add_paragraph(i['option']+". "+i['content']+"\n")

        document.add_paragraph(
            copyright_content)

        document.save("./"+contest_id+"/Documents/"+'Answer.docx')

        print("The Exam Paper had been Generated in " +
              '"./'+contest_id+'/Documents/"')

    else:
        print("比赛ID "+contest_id+"文件夹不存在")